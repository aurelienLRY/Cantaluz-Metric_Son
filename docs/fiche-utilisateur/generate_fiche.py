#!/usr/bin/env python3
"""Génère FICHE_UTILISATEUR.docx — guide utilisateur Cantaluz."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
IMAGES = ROOT / "images"


def set_a4_margins(section):
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


def shade_paragraph(paragraph, fill="FFF3CD"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def add_warning_box(doc, title: str, lines: list[str], fill="FEE2E2"):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
    shade_paragraph(p, fill)
    for line in lines:
        bp = doc.add_paragraph(line, style="List Bullet")
        shade_paragraph(bp, fill)
        for run in bp.runs:
            run.font.size = Pt(10.5)
    doc.add_paragraph()


def add_info_box(doc, title: str, text: str, fill="EFF6FF"):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    shade_paragraph(p, fill)
    bp = doc.add_paragraph(text)
    shade_paragraph(bp, fill)
    bp.runs[0].font.size = Pt(10.5)
    doc.add_paragraph()


def add_placeholder(doc, filename: str, caption: str, width_cm: float = 16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = IMAGES / filename
    if img_path.exists():
        run = p.add_run()
        run.add_picture(str(img_path), width=Cm(width_cm))
    else:
        run = p.add_run(f"[ Insérer une photo : {caption} ]")
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.size = Pt(11)
        run.italic = True
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    for r in p.runs:
        r.font.size = Pt(10.5)


def add_numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    p.runs[0].font.size = Pt(10.5)


def build():
    doc = Document()
    set_a4_margins(doc.sections[0])

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ── Page de titre ──────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Cantaluz")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)

    sub = doc.add_paragraph("Guide utilisateur")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(16)
    sub.runs[0].bold = True

    tag = doc.add_paragraph("La lumière qui répond au son")
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.runs[0].font.size = Pt(12)
    tag.runs[0].italic = True

    doc.add_paragraph()

    # ── 0. Avertissement (impératif) ───────────────────────────────
    doc.add_heading("Avertissement — à lire avant toute utilisation", level=1)
    add_warning_box(
        doc,
        "Règle impérative",
        [
            "Débrancher l'appareil (retirer l'alimentation) dès qu'il n'est plus utilisé "
            "(fin de séance, pause prolongée, vacances scolaires).",
            "Ne pas laisser Cantaluz branché en permanence sur une prise ou une batterie "
            "lorsqu'aucun adulte référent n'est présent.",
        ],
    )
    doc.add_paragraph(
        "Cantaluz est un outil électrique destiné à un usage scolaire encadré. "
        "Le non-respect de cette consigne expose à un risque de surchauffe, "
        "d'usure prématurée du matériel et de consommation inutile d'énergie."
    )

    # ── Plan du guide ──────────────────────────────────────────────
    doc.add_heading("Comment utiliser ce guide", level=1)
    doc.add_paragraph(
        "Ce document est pensé pour vous accompagner pas à pas, du premier branchement "
        "jusqu'à l'utilisation autonome en classe. Suivez les sections dans l'ordre "
        "lors de votre première prise en main ; ensuite, consultez directement la partie "
        "qui correspond à votre besoin."
    )
    add_table(
        doc,
        ["Étape", "Section", "Objectif"],
        [
            ("1", "C'est quoi Cantaluz ?", "Comprendre l'outil en 1 minute"),
            ("2", "Première utilisation", "Brancher, allumer — Cantaluz fonctionne tout de suite"),
            ("3", "Lire le ruban", "Savoir ce que signifie la lumière"),
            ("4", "Choisir un mode", "Changer de mode (nécessite le téléphone)"),
            ("5", "Ajuster si besoin", "Modifier les réglages (nécessite le téléphone)"),
            ("6", "Fin de séance", "Éteindre et débrancher correctement"),
            ("7", "Un problème ?", "Solutions rapides"),
        ],
        col_widths=[1.5, 5.5, 9],
    )

    # ── 1. C'est quoi Cantaluz ? ───────────────────────────────────
    doc.add_heading("1. C'est quoi Cantaluz ?", level=1)
    doc.add_paragraph(
        "Cantaluz est un ruban LED qui réagit au son de la classe. Un micro mesure "
        "l'ambiance ; le ruban affiche une barre lumineuse dont la hauteur et les couleurs "
        "indiquent le niveau sonore."
    )
    add_bullet(doc, "Peu de LED allumées = classe calme.", "En bref : ")
    add_bullet(doc, "Beaucoup de LED = classe bruyante.", "")
    add_bullet(doc, "Couleurs du vert à l'orange puis au rouge le long du ruban.", "")
    add_info_box(
        doc,
        "Le téléphone est optionnel",
        "Cantaluz fonctionne de façon autonome dès l'allumage : le ruban réagit au son "
        "en mode Flash, sans aucune configuration. Brancher, allumer, c'est prêt. "
        "Le téléphone n'est utile que pour changer de mode (Standard, Méditation, Défi Fifou) "
        "ou ajuster les réglages via l'application.",
    )
    add_placeholder(doc, "01-appareil.jpg", "Votre appareil Cantaluz (boîtier + ruban LED)", 14)

    # ── 2. Première utilisation ────────────────────────────────────
    doc.add_heading("2. Première utilisation — prêt en quelques secondes", level=1)
    doc.add_paragraph(
        "Cantaluz est utilisable immédiatement après l'allumage. "
        "Les étapes ci-dessous suffisent pour une séance en mode Flash, "
        "sans téléphone ni connexion Wi-Fi."
    )

    doc.add_heading("2.1 Brancher et allumer", level=2)
    add_numbered(doc, "Brancher l'alimentation 5 V (ruban LED + carte, masse commune).")
    add_numbered(doc, "Appuyer sur le bouton marche / arrêt.")
    add_numbered(doc, "Attendre environ 10 secondes sans rien toucher.")

    doc.add_heading("2.2 Ce qui se passe tout seul au démarrage", level=2)
    add_table(
        doc,
        ["Ordre", "Ce que vous voyez", "Signification"],
        [
            ("1", "Ruban bleu un instant", "L'appareil démarre"),
            ("2", "Animation verte qui monte puis redescend", "Test du ruban"),
            ("3", "Quelques LED vertes, le ruban réagit au son", "Mode d'écoute actif"),
            ("4", "LED bleue sur la carte", "Wi-Fi prêt"),
        ],
        col_widths=[1.8, 5.5, 8.7],
    )
    doc.add_paragraph(
        "À la fin de cette séquence, Cantaluz est en mode Flash (mode par défaut). "
        "Le ruban écoute la classe et réagit au son : vous pouvez commencer votre séance "
        "sans rien faire d'autre."
    )
    add_info_box(
        doc,
        "C'est tout pour une utilisation simple",
        "Si le mode Flash vous convient et que les réglages par défaut vous satisfont, "
        "inutile de sortir votre téléphone. Passez directement à la section 3 pour lire le ruban, "
        "ou à la section 6 en fin de séance.",
    )

    doc.add_heading("2.3 (Optionnel) Se connecter à l'application", level=2)
    doc.add_paragraph(
        "Cette étape n'est nécessaire que pour changer de mode ou modifier les réglages. "
        "Connectez votre téléphone au Wi-Fi de Cantaluz, puis ouvrez l'application dans le navigateur."
    )
    add_table(
        doc,
        ["", ""],
        [("Réseau Wi-Fi", "Cantaluz"), ("Mot de passe", "cantaluz1")],
        col_widths=[5, 11],
    )
    add_numbered(doc, "Sur le téléphone, ouvrir les réglages Wi-Fi et se connecter au réseau Cantaluz.")
    add_numbered(
        doc,
        "Si un message « pas d'Internet » s'affiche : rester connecté au Wi-Fi Cantaluz "
        "(ne pas repasser en 4G / données mobiles).",
    )
    add_numbered(
        doc,
        "Ouvrir Chrome et saisir http://cantaluz.local ou http://192.168.4.1 "
        "(la page peut aussi s'ouvrir automatiquement).",
    )
    add_info_box(
        doc,
        "Depuis l'application, vous pouvez :",
        "Changer de mode · Lancer une méditation ou un défi · Ajuster les réglages (Enregistrer). "
        "Sans connexion à l'app, Cantaluz reste en mode Flash avec les réglages d'origine.",
    )
    add_placeholder(doc, "03-capture-app.jpg", "L'application sur téléphone (onglet Dashboard)", 11)

    # ── 3. Lire le ruban ───────────────────────────────────────────
    doc.add_heading("3. Lire le ruban — comprendre ce que vous voyez", level=1)
    doc.add_paragraph(
        "Dans les modes Flash et Standard, le ruban fonctionne comme un « thermomètre du bruit » :"
    )
    mono = doc.add_paragraph()
    mono.add_run(
        "[Début du ruban]  ████████░░░░░░░░░░░░░░░░  [Fin du ruban]\n"
        "                  vert    orange      rouge\n"
        "                  ↑ calme              ↑ bruyant"
    ).font.name = "Consolas"
    mono.runs[0].font.size = Pt(9.5)

    add_table(
        doc,
        ["Ce que vous voyez", "Ce que ça veut dire"],
        [
            ("Peu de LED, surtout vertes", "La classe est calme"),
            ("Barre qui monte, couleurs orange", "Ambiance animée"),
            ("Barre haute, couleurs rouges", "Niveau sonore élevé"),
            ("Flash bleu sur tout le ruban (mode Flash)", "Un palier sonore vient d'être dépassé"),
        ],
        col_widths=[6, 10],
    )
    add_placeholder(
        doc,
        "02-ruban-en-action.jpg",
        "Exemple : ruban en action (barre vert → orange → rouge)",
        16,
    )

    # ── 4. Choisir un mode ─────────────────────────────────────────
    doc.add_heading("4. Choisir un mode — quel mode pour quelle situation ?", level=1)
    doc.add_paragraph(
        "Au démarrage, Cantaluz est toujours en mode Flash : aucune action requise. "
        "Pour utiliser un autre mode, connectez le téléphone à l'application "
        "(section 2.3) : onglet Dashboard → carte Mode. Le changement est immédiat sur le ruban. "
        "Le mode choisi reste actif jusqu'à l'extinction de l'appareil (au redémarrage : Flash par défaut)."
    )

    doc.add_heading("Quel mode choisir ?", level=2)
    add_table(
        doc,
        ["Votre situation", "Mode recommandé"],
        [
            ("Séance normale, suivre le bruit et signaler les dépassements", "Flash"),
            ("Travail calme prolongé (lecture, écriture), sans flashs", "Standard"),
            ("Classe agitée, pause respiration collective", "Méditation guidée"),
            ("Retour au calme en jeu (« remplissez le ruban ! »)", "Défi Fifou"),
        ],
        col_widths=[8, 8],
    )

    modes = [
        (
            "Flash — le mode du quotidien (actif par défaut)",
            "Barre réactive au son, couleurs vert → orange → rouge. "
            "Flashs bleus à chaque montée de palier sonore.",
            "Actif automatiquement à l'allumage — téléphone non requis. "
            "Pour revenir en Flash après un autre mode : sélectionner Flash dans le Dashboard.",
        ),
        (
            "Standard — accompagnement doux",
            "Même barre et mêmes couleurs, montée/descente plus fluide, sans flash bleu.",
            "Nécessite l'application : sélectionner Standard dans le Dashboard.",
        ),
        (
            "Méditation guidée — respiration par la lumière",
            "Le micro est désactivé. Cyan = inspire · Ambre = retiens · Magenta = expire · "
            "Lavande = bloque l'air. LED une par une sur le ruban.",
            "Nécessite l'application : Méditation guidée → durée 2, 5 ou 10 min → Démarrer "
            "(compte à rebours 5 s, puis chrono). Bouton Arrêter pour interrompre.",
        ),
        (
            "Défi Fifou — jeu du calme",
            "20 LED cyan au départ. Calme = les LED avancent · Animé = statu quo · "
            "Intense = recul. Victoire = feu d'artifice · Défaite = clignotement rouge.",
            "Nécessite l'application : Défi Fifou → durée 2, 5 ou 10 min → Démarrer. "
            "Astuce : en 2 min, viser environ 1 min 20 de calme continu.",
        ),
    ]

    for title, comportement, lancer in modes:
        doc.add_heading(title, level=2)
        doc.add_paragraph(comportement)
        p = doc.add_paragraph()
        r = p.add_run(lancer)
        r.italic = True
        r.font.size = Pt(10)

    # ── 5. Ajuster si besoin ───────────────────────────────────────
    doc.add_heading("5. Ajuster si besoin — onglet Réglages (téléphone requis)", level=1)
    doc.add_paragraph(
        "Les réglages par défaut conviennent dans la plupart des cas — "
        "aucune connexion au téléphone n'est nécessaire pour une séance en mode Flash. "
        "Connectez-vous à l'application uniquement si le ruban réagit trop ou pas assez : "
        "onglet Réglages (grisé pendant une Méditation ou un Défi Fifou)."
    )
    add_table(
        doc,
        ["Réglage", "Quand le modifier", "Effet"],
        [
            ("Sensibilité", "Ruban bouge au moindre bruit de fond", "Plus bas = plus stable au silence"),
            ("Zone calme", "Trop de flashs ou de rouge", "Augmenter = rester plus longtemps en vert"),
            ("Zone animée", "Passage trop rapide vers le rouge", "Augmenter = orange/rouge plus tardifs"),
            ("Luminosité", "Ruban trop vif ou trop faible", "Ajuster le confort visuel"),
            ("Montée de la barre", "Réaction trop lente ou trop nerveuse", "Vitesse de montée de la barre"),
        ],
        col_widths=[3.5, 5, 7.5],
    )
    add_info_box(
        doc,
        "Important",
        "Cliquer sur Enregistrer pour appliquer les réglages. "
        "Ils reviennent aux valeurs d'origine au prochain redémarrage de l'appareil.",
    )

    # ── 6. Fin de séance ───────────────────────────────────────────
    doc.add_heading("6. Fin de séance — éteindre et ranger", level=1)
    doc.add_paragraph("À la fin de chaque utilisation, effectuez ces deux gestes dans l'ordre :")
    add_numbered(doc, "Appuyer sur le bouton marche / arrêt pour éteindre Cantaluz.")
    add_numbered(doc, "Débrancher l'alimentation (retirer le câble USB ou l'adaptateur secteur).")
    add_warning_box(
        doc,
        "Rappel",
        [
            "Ne jamais laisser Cantaluz branché lorsqu'il n'est pas utilisé.",
        ],
        fill="FEE2E2",
    )

    # ── 7. Dépannage ───────────────────────────────────────────────
    doc.add_heading("7. Un problème ? — solutions rapides", level=1)
    add_table(
        doc,
        ["Problème", "Que faire"],
        [
            ("Le ruban ne réagit pas au son", "Vérifier l'alimentation ; attendre la fin de la séquence de démarrage (~10 s)"),
            ("Pas de Wi-Fi Cantaluz (connexion app)", "Vérifier que l'appareil est branché et allumé ; attendre 10 s ; LED bleue sur la carte ?"),
            ("Page web ne s'ouvre pas", "Chrome → http://192.168.4.1 ; couper les données mobiles automatiques"),
            ("Ruban bouge tout seul", "Réglages → baisser Sensibilité (essayer entre 0 et 10)"),
            ("Trop de rouge / flashs", "Baisser Sensibilité ; augmenter Zone calme et Zone animée"),
            ("Méditation ne démarre pas", "Choisir d'abord le mode, puis la durée, puis Démarrer"),
            ("Défi Fifou ne démarre pas", "Même procédure : mode → durée → Démarrer"),
        ],
        col_widths=[5.5, 10.5],
    )

    # ── Rappels matériel ───────────────────────────────────────────
    doc.add_heading("Rappels matériel & sécurité", level=1)
    for item in [
        "Alimentation 5 V adaptée, masse commune entre ruban et carte.",
        "Ne pas regarder le ruban de très près en pleine luminosité prolongée.",
        "Usage scolaire avec adulte référent uniquement.",
        "Ne pas connecter l'ordinateur (moniteur série) en usage normal — cela peut gêner le Wi-Fi.",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph()
    foot = doc.add_paragraph("Cantaluz — guide utilisateur · format A4 · opteo-web.fr")
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.runs[0].font.size = Pt(9)
    foot.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    out = ROOT / "FICHE_UTILISATEUR.docx"
    doc.save(out)
    print(f"Généré : {out}")


if __name__ == "__main__":
    build()
